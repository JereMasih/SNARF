import io

from docx import Document

from snarf.capabilities.base import Capability


class DocxExtractor(Capability):
    name = "docx_extractor"

    @property
    def available(self) -> bool:
        return True

    def extract_text(self, docx_bytes: bytes) -> str:
        document = Document(io.BytesIO(docx_bytes))
        parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
