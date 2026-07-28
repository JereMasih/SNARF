import io

from docx import Document

from snarf.capabilities.docx_extractor import DocxExtractor


def make_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("primer párrafo real")
    document.add_paragraph("segundo párrafo real")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "celda A"
    table.rows[0].cells[1].text = "celda B"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_extract_text_includes_paragraphs():
    text = DocxExtractor().extract_text(make_docx_bytes())
    assert "primer párrafo real" in text
    assert "segundo párrafo real" in text


def test_extract_text_includes_table_content():
    text = DocxExtractor().extract_text(make_docx_bytes())
    assert "celda A" in text
    assert "celda B" in text


def test_is_always_available():
    assert DocxExtractor().available is True
